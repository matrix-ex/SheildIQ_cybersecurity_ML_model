import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        stripped_line = line.strip()

        # Handle Tables
        if stripped_line.startswith('|') and stripped_line.endswith('|'):
            if '---' in stripped_line: # Skip separator lines
                continue
            
            in_table = True
            # Parse row
            cells = [c.strip() for c in stripped_line.split('|') if c.strip() or stripped_line.count('|') > 1]
            if not cells and stripped_line == '||': # Edge case for empty cells
                cells = ['']
            table_data.append(cells)
            continue
        elif in_table:
            # End of table detected
            if table_data:
                # Filter out empty rows if any
                table_data = [row for row in table_data if any(row)]
                if table_data:
                    num_cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=0, cols=num_cols)
                    table.style = 'Table Grid'
                    for row_cells in table_data:
                        row = table.add_row()
                        for i, cell_text in enumerate(row_cells):
                            if i < num_cols:
                                row.cells[i].text = cell_text
            table_data = []
            in_table = False

        # Handle Headings
        if stripped_line.startswith('# '):
            p = doc.add_heading(stripped_line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('#### '):
            doc.add_heading(stripped_line[5:], level=4)
        
        # Handle Bullet Points
        elif stripped_line.startswith('- '):
            doc.add_paragraph(stripped_line[2:], style='List Bullet')
        
        # Handle Page Breaks (explicitly requested placeholders or section markers)
        elif '---' in stripped_line and not in_table:
            doc.add_page_break()

        # Handle Images (Placeholders)
        elif '[Insert Screenshot Here]' in stripped_line:
            p = doc.add_paragraph()
            run = p.add_run('[INSERT SCREENSHOT HERE]')
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Regular Text
        elif stripped_line:
            doc.add_paragraph(line.rstrip('\n')) # Keep original indentation for trees
        else:
            # Empty line
            pass

    # Final table check if file ends with a table
    if in_table and table_data:
        table_data = [row for row in table_data if any(row)]
        if table_data:
            num_cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = 'Table Grid'
            for row_cells in table_data:
                row = table.add_row()
                for i, cell_text in enumerate(row_cells):
                    if i < num_cols:
                        row.cells[i].text = cell_text

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx('VAULTO_Final_Report.md', 'VAULTO_Final_Report.docx')
